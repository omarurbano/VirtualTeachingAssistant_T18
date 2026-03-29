"""
Unified Document Processor using Gemini API
=============================================

This module provides comprehensive document processing using Google Gemini API:
- PDF and DOCX support with text, tables, and images
- Deep image analysis using Gemini Vision API
- Table extraction and analysis
- Audio transcription with timestamps
- All embeddings using Gemini Embedding 2.0

Author: CPT_S 421 Development Team
Version: 2.0.0
Created: 2026-03-28
"""

# ============================================================================
# STANDARD LIBRARY IMPORTS
# ============================================================================

import os
import io
import logging
import base64
import hashlib
import json
import re
from typing import List, Dict, Any, Optional, Tuple
from datetime import datetime
from dataclasses import dataclass, field

# ============================================================================
# THIRD-PARTY LIBRARY IMPORTS
# ============================================================================

# Google Gemini
try:
    import google.generativeai as genai
    GEMINI_AVAILABLE = True
except ImportError:
    GEMINI_AVAILABLE = False
    logging.warning("google-generativeai not installed")

# PDF processing
try:
    import PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False
    logging.warning("PyMuPDF not installed")

# Word documents
try:
    from docx import Document
    from docx.table import Table
    from docx.image.image import BaseImage
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx not installed")

# Image processing
try:
    from PIL import Image
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False
    logging.warning("PIL not installed")

# Numerical operations
try:
    import numpy as np
    NUMPY_AVAILABLE = True
except ImportError:
    NUMPY_AVAILABLE = False
    logging.warning("numpy not installed")

# ============================================================================
# LOGGING CONFIGURATION
# ============================================================================

logger = logging.getLogger(__name__)

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)

# ============================================================================
# CONFIGURATION
# ============================================================================

# Gemini models
GEMINI_VISION_MODEL = 'gemini-2.0-flash'  # Vision model for images
GEMINI_AUDIO_MODEL = 'gemini-2.0-flash'  # For audio transcription
GEMINI_ANALYSIS_MODEL = 'gemini-2.0-flash'  # For table/text analysis

# Embedding model
GEMINI_EMBEDDING_MODEL = 'models/gemini-embedding-2-preview'
GEMINI_EMBEDDING_DIMENSIONS = 768

# ============================================================================
# DATA CLASSES
# ============================================================================

@dataclass
class DocumentChunk:
    """Represents a chunk from any document type."""
    content: str
    chunk_type: str  # 'text', 'table', 'image', 'audio'
    source_file: str
    document_id: str
    page_number: Optional[int] = None
    timestamp: Optional[float] = None  # For audio
    image_data: Optional[str] = None  # Base64 encoded image
    table_data: Optional[str] = None  # Markdown table
    embedding: Optional[List[float]] = None
    metadata: Dict[str, Any] = field(default_factory=dict)
    
    def to_dict(self) -> Dict:
        """Convert to dictionary."""
        return {
            'content': self.content,
            'chunk_type': self.chunk_type,
            'source_file': self.source_file,
            'document_id': self.document_id,
            'page_number': self.page_number,
            'timestamp': self.timestamp,
            'metadata': self.metadata
        }


@dataclass
class ProcessedDocument:
    """Represents a fully processed document."""
    document_id: str
    filename: str
    file_type: str  # 'pdf', 'docx', 'audio'
    chunks: List[DocumentChunk]
    images: List[Dict[str, Any]] = field(default_factory=list)
    tables: List[Dict[str, Any]] = field(default_factory=list)
    transcript: Optional[str] = None
    processing_time: float = 0.0
    metadata: Dict[str, Any] = field(default_factory=dict)


# ============================================================================
# GEMINI CLIENT
# ============================================================================

class GeminiClient:
    """
    Unified Gemini API client for all operations.
    """
    
    def __init__(self, api_key: str = None):
        """
        Initialize Gemini client.
        
        Args:
            api_key: Google API key (falls back to GOOGLE_API_KEY env var)
        """
        if not GEMINI_AVAILABLE:
            raise ImportError("google-generativeai not installed")
        
        if api_key is None:
            api_key = os.environ.get('GOOGLE_API_KEY')
        
        if not api_key:
            raise ValueError("API key required. Set GOOGLE_API_KEY or pass as parameter")
        
        genai.configure(api_key=api_key)
        self.api_key = api_key
        self.embedding_model = GEMINI_EMBEDDING_MODEL
        self.dimensions = GEMINI_EMBEDDING_DIMENSIONS
        
        logger.info("Gemini client initialized")
    
    # ==================== TEXT EMBEDDINGS ====================
    
    def embed_texts(self, texts: List[str], task_type: str = 'retrieval_document') -> List[List[float]]:
        """
        Generate embeddings for texts using Gemini Embedding 2.0.
        
        Args:
            texts: List of text strings
            task_type: Task optimization type
            
        Returns:
            List of embedding vectors
        """
        if not texts:
            return []
        
        try:
            result = genai.embed_content(
                model=self.embedding_model,
                content=texts,
                task_type=task_type,
                output_dimensionality=self.dimensions
            )
            return result['embedding']
        except Exception as e:
            logger.error(f"Embedding error: {e}")
            raise
    
    def embed_query(self, query: str) -> List[float]:
        """
        Generate embedding for a query.
        
        Args:
            query: Query string
            
        Returns:
            Query embedding vector
        """
        try:
            result = genai.embed_content(
                model=self.embedding_model,
                content=query,
                task_type='retrieval_query',
                output_dimensionality=self.dimensions
            )
            return result['embedding'][0]
        except Exception as e:
            logger.error(f"Query embedding error: {e}")
            raise
    
    def embed_image(self, image_data: bytes) -> List[float]:
        """
        Generate embedding for an image.
        
        Args:
            image_data: Image bytes
            
        Returns:
            Image embedding vector
        """
        try:
            # Convert to base64
            b64_image = base64.b64encode(image_data).decode('utf-8')
            
            image_content = {
                'mime_type': 'image/png',
                'data': b64_image
            }
            
            result = genai.embed_content(
                model=self.embedding_model,
                content=[image_content],
                task_type='retrieval_document',
                output_dimensionality=self.dimensions
            )
            return result['embedding'][0]
        except Exception as e:
            logger.error(f"Image embedding error: {e}")
            raise
    
    # ==================== IMAGE ANALYSIS ====================
    
    def analyze_image(self, image_data: bytes, prompt: str = None) -> Dict[str, Any]:
        """
        Analyze image using Gemini Vision API.
        
        Args:
            image_data: Image bytes
            prompt: Optional prompt for analysis
            
        Returns:
            Dict with analysis results
        """
        try:
            # Upload image to Gemini
            image_io = io.BytesIO(image_data)
            
            model = genai.GenerativeModel(GEMINI_VISION_MODEL)
            
            # Default prompt for comprehensive analysis
            if prompt is None:
                prompt = """Analyze this image thoroughly. Provide:
1. A detailed description of what's in the image
2. Any text visible in the image
3. Any charts, graphs, or data visualizations
4. The context or subject matter
5. Key details that would help find this image by searching
"""
            
            # Generate content
            response = model.generate_content([
                prompt,
                {'mime_type': 'image/png', 'data': base64.b64encode(image_data).decode('utf-8')}
            ])
            
            # Parse response
            analysis = {
                'description': response.text,
                'success': True,
                'model': GEMINI_VISION_MODEL
            }
            
            # Try to extract structured info
            return analysis
            
        except Exception as e:
            logger.error(f"Image analysis error: {e}")
            return {
                'description': '',
                'success': False,
                'error': str(e)
            }
    
    def analyze_image_with_vision(self, image_data: bytes) -> Dict[str, Any]:
        """
        Deep image analysis using Gemini's full vision capabilities.
        
        Args:
            image_data: Image bytes
            
        Returns:
            Comprehensive analysis dict
        """
        try:
            model = genai.GenerativeModel(GEMINI_VISION_MODEL)
            
            # Enhanced detailed analysis prompt for chart/graph interpretation, OCR confidence, and object detection
            prompt = """You are an expert image analyst with deep understanding of charts, graphs, technical diagrams, and document analysis. Provide a comprehensive analysis in JSON format:

{
  "description": "Detailed description of all visual elements, layout, and composition",
  "text_content": {
    "raw_text": "All text visible in the image (transcribed exactly)",
    "text_blocks": [
      {
        "text": "text string",
        "confidence": 0.95,  // OCR confidence score (0-1)
        "bounding_box": [x1, y1, x2, y2],  // Normalized coordinates (0-1) as [left, top, right, bottom]
        "font_size_estimate": "small/medium/large",
        "is_header": true/false
      }
    ]
  },
  "data_visualization": {
    "chart_type": "bar/line/pie/scatter/area/histogram/boxplot/unknown",
    "data_points": [
      {
        "label": "category or x-value",
        "value": numeric_value,
        "series": "series name if applicable"
      }
    ],
    "axes": {
      "x_axis": {
        "label": "axis label",
        "type": "categorical/numeric/date",
        "range": [min, max]  // for numeric axes
      },
      "y_axis": {
        "label": "axis label", 
        "type": "categorical/numeric/date",
        "range": [min, max]  // for numeric axes
      }
    },
    "legends": [
      {
        "title": "legend title",
        "items": ["item1", "item2"]
      }
    ],
    "title": "chart or graph title",
    "caption": "caption or description below chart",
    "trends": ["increasing", "decreasing", "peak at X", "outlier at Y"],
    "statistics_visible": ["mean: 5.2", "std: 1.3"]  // Any statistics shown in chart
  },
  "objects_detected": [
    {
      "object": "person/chair/computer/etc",
      "confidence": 0.9,
      "bounding_box": [x1, y1, x2, y2],
      "description": "brief description of the object and its context"
    }
  ],
  "context": {
    "document_type": "research paper/presentation slide/report/dashboard/other",
    "subject_area": "finance/engineering/medicine/etc if detectable",
    "likely_purpose": "what this image is likely used for",
    "page_layout": "description of overall layout"
  },
  "search_keywords": ["keyword1", "keyword2", "keyword3", "keyword4", "keyword5"],
  "confidence_scores": {
    "overall": 0.0-1.0,
    "chart_interpretation": 0.0-1.0,
    "text_extraction": 0.0-1.0,
    "object_detection": 0.0-1.0
  }
}"""
            
            response = model.generate_content([
                prompt,
                {'mime_type': 'image/png', 'data': base64.b64encode(image_data).decode('utf-8')}
            ])
            
            # Try to parse as JSON, fallback to plain text
            try:
                # Find JSON in response
                json_match = re.search(r'\{.*\}', response.text, re.DOTALL)
                if json_match:
                    analysis = json.loads(json_match.group())
                    
                    # Validate and enhance the analysis structure
                    if 'text_content' not in analysis:
                        analysis['text_content'] = {
                            'raw_text': analysis.get('description', ''),
                            'text_blocks': []
                        }
                        
                    if 'data_visualization' not in analysis:
                        analysis['data_visualization'] = {
                            'chart_type': 'unknown',
                            'data_points': [],
                            'axes': {},
                            'legends': [],
                            'title': '',
                            'caption': '',
                            'trends': [],
                            'statistics_visible': []
                        }
                        
                    if 'objects_detected' not in analysis:
                        analysis['objects_detected'] = []
                        
                    if 'context' not in analysis:
                        analysis['context'] = {
                            'document_type': 'unknown',
                            'subject_area': 'unknown',
                            'likely_purpose': 'unknown',
                            'page_layout': 'unknown'
                        }
                    if 'search_keywords' not in analysis:
                        analysis['search_keywords'] = []
                    if 'confidence_scores' not in analysis:
                        analysis['confidence_scores'] = {
                            'overall': 0.5,
                            'chart_interpretation': 0.5,
                            'text_extraction': 0.5,
                            'object_detection': 0.5
                        }
                else:
                    analysis = {
                        'description': response.text, 
                        'text_content': {'raw_text': '', 'text_blocks': []},
                        'data_visualization': {'chart_type': 'unknown', 'data_points': [], 'axes': {}, 'legends': [], 'title': '', 'caption': '', 'trends': [], 'statistics_visible': []},
                        'objects_detected': [],
                        'context': {'document_type': 'unknown', 'subject_area': 'unknown', 'likely_purpose': 'unknown', 'page_layout': 'unknown'},
                        'search_keywords': [],
                        'confidence_scores': {
                            'overall': 0.5,
                            'chart_interpretation': 0.5,
                            'text_extraction': 0.5,
                            'object_detection': 0.5
                        }
                    }
            except Exception as json_error:
                logger.warning(f"JSON parsing failed in image analysis: {json_error}")
                analysis = {
                    'description': response.text, 
                    'text_content': {'raw_text': '', 'text_blocks': []},
                    'data_visualization': {'chart_type': 'unknown', 'data_points': [], 'axes': {}, 'legends': [], 'title': '', 'caption': '', 'trends': [], 'statistics_visible': []},
                    'objects_detected': [],
                    'context': {'document_type': 'unknown', 'subject_area': 'unknown', 'likely_purpose': 'unknown', 'page_layout': 'unknown'},
                    'search_keywords': [],
                    'confidence_scores': {
                        'overall': 0.5,
                        'chart_interpretation': 0.5,
                        'text_extraction': 0.5,
                        'object_detection': 0.5
                    }
                }
            
            analysis['success'] = True
            analysis['raw_response'] = response.text
            
            return analysis
            
        except Exception as e:
            logger.error(f"Vision analysis error: {e}")
            return {'success': False, 'error': str(e), 'description': ''}
    
    # ==================== TABLE ANALYSIS ====================
    
    def analyze_table(self, table_markdown: str, context: str = '') -> Dict[str, Any]:
        """
        Analyze a table using Gemini with detailed structural analysis.
        
        Args:
            table_markdown: Table in markdown format
            context: Optional context about the table
            
        Returns:
            Detailed table analysis results with structure, columns, relationships
        """
        try:
            model = genai.GenerativeModel(GEMINI_ANALYSIS_MODEL)
            
            # Count rows and columns from markdown
            rows = table_markdown.strip().split('\n')
            row_count = len([r for r in rows if r.strip() and '|' in r])
            col_count = 0
            if rows:
                first_row = rows[0]
                col_count = len([c for c in first_row.split('|') if c.strip()])
            
            prompt = f"""You are an expert data analyst specializing in table interpretation. Analyze this table comprehensively and provide a detailed analysis in JSON format:

{{
  "summary": "Concise summary of what the table represents and its main purpose",
  "structure": {{
    "row_count": {row_count},
    "column_count": {col_count},
    "has_header_row": true/false,
    "has_footer_row": true/false,
    "is_numeric_matrix": true/false,
    "table_type": "financial/scientific/survey/schedule/inventory/other"
  }},
  "columns": [
    {{
      "column_index": 0,
      "header": "column header text",
      "data_type": "numeric/integer/float/currency/percentage/date/time/text/categorical/mixed",
      "format_detected": "e.g., $1,234.56, 25%, 2023-01-15, etc.",
      "is_key_column": true/false,
      "is_foreign_key": true/false,
      "unique_value_count": estimated_count_or_null,
      "null_or_empty_count": count_of_empty_cells,
      "descriptive_statistics": {{
        "mean": numeric_value_or_null,
        "median": numeric_value_or_null,
        "mode": "most_frequent_value_or_null",
        "std_dev": numeric_value_or_null,
        "min": numeric_value_or_null,
        "max": numeric_value_or_null
      }},
      "value_range": {{"min": min_value, "max": max_value}},
      "most_common_values": [{{"value": "value1", "count": count1}}],
      "data_quality_issues": ["missing_values", "inconsistent_format", "outliers"],
      "semantic_type": "id/name/description/quantity/price/percentage/rate/score/etc"
    }}
  ],
  "relationships": [
    {{
      "type": "functional_dependency/correlation/categorical_grouping",
      "description": "description of the relationship",
      "involved_columns": [0, 1, 2],
      "strength": "strong/medium/weak"
    }}
  ],
  "patterns_insights": [
    "Specific observations about data patterns",
    "Trends visible in the data",
    "Anomalies or outliers detected"
  ],
  "data_quality": {{
    "overall_score": 0.0-1.0,
    "completeness": 0.0-1.0,
    "consistency": 0.0-1.0,
    "accuracy": 0.0-1.0,
    "issues_found": [
      {{"type": "missing_values", "location": "column 2, rows 5-7", "severity": "medium"}}
    ]
  }},
  "suggested_visualizations": ["bar_chart", "line_chart", "pie_chart", "scatter_plot", "heatmap"],
  "potential_use_cases": ["financial_reporting", "inventory_tracking", "survey_analysis"],
  "confidence_scores": {{
    "overall": 0.0-1.0,
    "structure_detection": 0.0-1.0,
    "type_detection": 0.0-1.0,
    "relationship_analysis": 0.0-1.0
  }}
}}

Table:
{table_markdown}

Context: {context}

Return valid JSON only. Do not include any text outside the JSON object."""
            
            response = model.generate_content(prompt)
            
            # Try to parse as JSON, fallback to plain text
            try:
                json_match = re.search(r'\{.*\}', response.text, re.DOTALL)
                if json_match:
                    analysis = json.loads(json_match.group())
                    
                    # Validate and ensure required fields exist with defaults
                    if 'summary' not in analysis:
                        analysis['summary'] = response.text[:200] + ("..." if len(response.text) > 200 else "")
                    if 'structure' not in analysis:
                        analysis['structure'] = {
                            'row_count': row_count,
                            'column_count': col_count,
                            'has_header_row': True,
                            'has_footer_row': False,
                            'is_numeric_matrix': False,
                            'table_type': 'unknown'
                        }
                    if 'columns' not in analysis:
                        analysis['columns'] = []
                    if 'relationships' not in analysis:
                        analysis['relationships'] = []
                    if 'patterns_insights' not in analysis:
                        analysis['patterns_insights'] = []
                    if 'data_quality' not in analysis:
                        analysis['data_quality'] = {
                            'overall_score': 0.5,
                            'completeness': 0.5,
                            'consistency': 0.5,
                            'accuracy': 0.5,
                            'issues_found': []
                        }
                    if 'suggested_visualizations' not in analysis:
                        analysis['suggested_visualizations'] = []
                    if 'potential_use_cases' not in analysis:
                        analysis['potential_use_cases'] = []
                    if 'confidence_scores' not in analysis:
                        analysis['confidence_scores'] = {
                            'overall': 0.5,
                            'structure_detection': 0.5,
                            'type_detection': 0.5,
                            'relationship_analysis': 0.5
                        }
                else:
                    # Fallback to basic analysis if no JSON found
                    analysis = {
                        'summary': response.text,
                        'structure': {
                            'row_count': row_count,
                            'column_count': col_count,
                            'has_header_row': True,
                            'has_footer_row': False,
                            'is_numeric_matrix': False,
                            'table_type': 'unknown'
                        },
                        'columns': [],
                        'relationships': [],
                        'patterns_insights': [],
                        'data_quality': {
                            'overall_score': 0.5,
                            'completeness': 0.5,
                            'consistency': 0.5,
                            'accuracy': 0.5,
                            'issues_found': []
                        },
                        'suggested_visualizations': [],
                        'potential_use_cases': [],
                        'confidence_scores': {
                            'overall': 0.5,
                            'structure_detection': 0.5,
                            'type_detection': 0.5,
                            'relationship_analysis': 0.5
                        }
                    }
            except Exception as json_error:
                logger.warning(f"Table analysis JSON parsing error: {json_error}")
                analysis = {
                    'summary': response.text,
                    'structure': {
                        'row_count': row_count,
                        'column_count': col_count,
                        'has_header_row': True,
                        'has_footer_row': False,
                        'is_numeric_matrix': False,
                        'table_type': 'unknown'
                    },
                    'columns': [],
                    'relationships': [],
                    'patterns_insights': [],
                    'data_quality': {
                        'overall_score': 0.5,
                        'completeness': 0.5,
                        'consistency': 0.5,
                        'accuracy': 0.5,
                        'issues_found': []
                    },
                    'suggested_visualizations': [],
                    'potential_use_cases': [],
                    'confidence_scores': {
                        'overall': 0.5,
                        'structure_detection': 0.5,
                        'type_detection': 0.5,
                        'relationship_analysis': 0.5
                    }
                }
            
            analysis['success'] = True
            analysis['table_markdown'] = table_markdown
            
            return analysis
            
        except Exception as e:
            logger.error(f"Table analysis error: {e}")
            return {
                'summary': '',
                'success': False,
                'error': str(e),
                'table_markdown': table_markdown,
                'structure': {'row_count': 0, 'column_count': 0, 'has_header_row': False, 'has_footer_row': False, 'is_numeric_matrix': False, 'table_type': 'unknown'},
                'columns': [],
                'relationships': [],
                'patterns_insights': [],
                'data_quality': {'overall_score': 0.0, 'completeness': 0.0, 'consistency': 0.0, 'accuracy': 0.0, 'issues_found': []},
                'suggested_visualizations': [],
                'potential_use_cases': [],
                'confidence_scores': {'overall': 0.0, 'structure_detection': 0.0, 'type_detection': 0.0, 'relationship_analysis': 0.0}
            }
    
    # ==================== AUDIO TRANSCRIPTION ====================
    
    def transcribe_audio(self, audio_data: bytes, mime_type: str = 'audio/mp3') -> Dict[str, Any]:
        """
        Transcribe audio using Gemini.
        
        Args:
            audio_data: Audio file bytes
            mime_type: MIME type of audio
            
        Returns:
            Transcription results with timestamps
        """
        import tempfile
        
        try:
            # Save bytes to a temporary file for Gemini upload
            # (genai.upload_file requires a file path, not bytes)
            ext = mime_type.split('/')[-1]
            if ext == 'mp4':
                ext = 'm4a'
            temp_file = tempfile.NamedTemporaryFile(suffix=f'.{ext}', delete=False)
            temp_file.write(audio_data)
            temp_file.flush()
            temp_file.close()
            
            logger.info(f"Uploading audio to Gemini: {len(audio_data)} bytes, mime={mime_type}")
            
            # Upload to Gemini using file path
            audio_file = genai.upload_file(
                path=temp_file.name,
                mime_type=mime_type
            )
            
            # Use Gemini 2.0 flash for transcription
            model = genai.GenerativeModel(GEMINI_AUDIO_MODEL)
            
            # Enhanced prompt for speaker diarization and speech characteristics
            prompt = """You are an expert audio transcription specialist. Provide a detailed transcription with speaker identification.

For each distinct speaker segment, provide:
1. The speaker identifier (Speaker 1, Speaker 2, etc.)
2. The timestamp (start time) in format [MM:SS] or [HH:MM:SS]
3. The transcribed text for that segment
4. Speech characteristics: tone/emotion (neutral, happy, sad, angry, excited, professional, calm, frustrated, etc.)
5. Confidence score for the transcription (0-1)

Format your response as a JSON array of segments:
[
  {
    "speaker": "Speaker 1",
    "timestamp": "[00:05]",
    "text": "Hello everyone, welcome to the presentation.",
    "tone": "neutral",
    "confidence": 0.92
  },
  {
    "speaker": "Speaker 2",
    "timestamp": "[00:15]",
    "text": "Thank you, I'll be discussing the quarterly results.",
    "tone": "professional",
    "confidence": 0.89
  }
]

Additionally, provide an overall summary at the end as JSON:
{
  "audio_quality": "clear/muffled/distorted/noisy/clean/poor/excellent/good/fair",
  "background_noise": "none/static/hiss/hum/echo/reverb/music/description",
  "language_detected": "english/spanish/french/etc",
  "estimated_speakers": 2
}

Important: Return valid JSON only. Wrap the segments array and summary in a single JSON object like:
{
  "segments": [...],
  "summary": {...}
}
"""
            
            response = model.generate_content([
                prompt,
                audio_file
            ])
            
            # Parse enhanced transcript with speaker diarization
            transcript_segments = self._parse_transcript(response.text)
            
            # Extract additional audio analysis from response
            audio_quality = self._extract_audio_quality(response.text)
            background_noise = self._extract_background_noise(response.text)
            language_detected = self._extract_language(response.text)
            estimated_speakers = len(set(seg.get('speaker', 'Unknown') for seg in transcript_segments))
            speech_rate_wpm = self._calculate_speech_rate(transcript_segments)
            
            # Clean up uploaded file and temp file
            try:
                genai.delete_file(audio_file.name)
            except:
                pass
            try:
                os.unlink(temp_file.name)
            except:
                pass
            
            return {
                'transcript': response.text,
                'segments': transcript_segments,
                'success': True,
                'full_text': '\n'.join([s['text'] for s in transcript_segments]),
                # Enhanced analysis fields
                'audio_quality': audio_quality,
                'background_noise': background_noise,
                'language_detected': language_detected,
                'estimated_speakers': estimated_speakers,
                'speech_rate_wpm': speech_rate_wpm
            }
            
        except Exception as e:
            logger.error(f"Audio transcription error: {e}")
            # Clean up temp file on error
            try:
                if 'temp_file' in dir():
                    os.unlink(temp_file.name)
            except:
                pass
            return {
                'transcript': '',
                'segments': [],
                'success': False,
                'error': str(e)
            }
    
    def _parse_transcript(self, transcript: str) -> List[Dict]:
        """
        Parse transcript text to extract speaker diarization and timestamps.
        
        Handles both JSON format (with speaker info) and legacy timestamp format.
        
        Args:
            transcript: Raw transcript text (expected JSON format)
            
        Returns:
            List of segment dicts with speaker, timestamp, text, tone, confidence
        """
        segments = []
        
        try:
            # Try to parse as JSON first (enhanced format)
            json_match = re.search(r'\{.*\}', transcript, re.DOTALL)
            if json_match:
                parsed = json.loads(json_match.group())
                
                # Handle wrapped format with "segments" key
                if 'segments' in parsed:
                    parsed_segments = parsed['segments']
                elif isinstance(parsed, list):
                    parsed_segments = parsed
                else:
                    parsed_segments = []
                
                for seg in parsed_segments:
                    if isinstance(seg, dict):
                        timestamp_str = seg.get('timestamp', '[00:00]').strip('[]')
                        timestamp_seconds = self._parse_timestamp_to_seconds(timestamp_str)
                        
                        segments.append({
                            'speaker': seg.get('speaker', 'Unknown'),
                            'timestamp': timestamp_str,
                            'timestamp_seconds': timestamp_seconds,
                            'text': seg.get('text', '').strip(),
                            'tone': seg.get('tone', 'neutral'),
                            'confidence': float(seg.get('confidence', 0.8)),
                            'word_count': len(seg.get('text', '').split())
                        })
            
            # If no JSON found, try array format
            if not segments:
                array_match = re.search(r'\[.*\]', transcript, re.DOTALL)
                if array_match:
                    try:
                        parsed_segments = json.loads(array_match.group())
                        for seg in parsed_segments:
                            if isinstance(seg, dict):
                                timestamp_str = seg.get('timestamp', '[00:00]').strip('[]')
                                timestamp_seconds = self._parse_timestamp_to_seconds(timestamp_str)
                                
                                segments.append({
                                    'speaker': seg.get('speaker', 'Unknown'),
                                    'timestamp': timestamp_str,
                                    'timestamp_seconds': timestamp_seconds,
                                    'text': seg.get('text', '').strip(),
                                    'tone': seg.get('tone', 'neutral'),
                                    'confidence': float(seg.get('confidence', 0.8)),
                                    'word_count': len(seg.get('text', '').split())
                                })
                    except json.JSONDecodeError:
                        pass
            
            # Fallback to original regex parsing for non-JSON responses
            if not segments:
                pattern = r'(?:\[?(\d{1,2}:\d{2})(?::\d{2})?\]?)\s*(.+?)(?=(?:\[?\d{1,2}:\d{2})|$)'
                matches = re.findall(pattern, transcript, re.DOTALL)
                
                for timestamp, text in matches:
                    segments.append({
                        'speaker': 'Speaker 1',  # Default when no speaker info
                        'timestamp': timestamp,
                        'timestamp_seconds': self._parse_timestamp_to_seconds(timestamp),
                        'text': text.strip(),
                        'tone': 'neutral',
                        'confidence': 0.8,  # Default confidence
                        'word_count': len(text.split())
                    })
                    
        except Exception as e:
            logger.error(f"Transcript parsing error: {e}")
        
        # Ultimate fallback: treat as single segment
        if not segments:
            segments.append({
                'speaker': 'Speaker 1',
                'timestamp': '00:00',
                'timestamp_seconds': 0.0,
                'text': transcript.strip(),
                'tone': 'neutral',
                'confidence': 0.5,
                'word_count': len(transcript.split())
            })
        
        return segments
    
    def _parse_timestamp_to_seconds(self, timestamp_str: str) -> float:
        """
        Convert timestamp string like 'MM:SS' or 'HH:MM:SS' to seconds.
        
        Args:
            timestamp_str: Timestamp string (e.g., '01:30' or '01:30:45')
            
        Returns:
            Float representing seconds
        """
        try:
            parts = timestamp_str.split(':')
            if len(parts) == 2:
                return int(parts[0]) * 60 + int(parts[1])
            elif len(parts) == 3:
                return int(parts[0]) * 3600 + int(parts[1]) * 60 + int(parts[2])
            else:
                return 0.0
        except (ValueError, IndexError):
            return 0.0
    
    def _extract_audio_quality(self, response_text: str) -> str:
        """
        Extract audio quality assessment from Gemini response.
        
        Args:
            response_text: Raw response text from Gemini
            
        Returns:
            Quality assessment string
        """
        quality_indicators = ['excellent', 'good', 'clear', 'clean', 'fair', 'poor', 'muffled', 'distorted', 'noisy']
        response_lower = response_text.lower()
        
        # Check for JSON-formatted quality info first
        try:
            json_match = re.search(r'"audio_quality"\s*:\s*"([^"]+)"', response_lower)
            if json_match:
                return json_match.group(1)
        except:
            pass
        
        for indicator in quality_indicators:
            if indicator in response_lower:
                return indicator
        return 'unknown'
    
    def _extract_background_noise(self, response_text: str) -> str:
        """
        Extract background noise description from Gemini response.
        
        Args:
            response_text: Raw response text from Gemini
            
        Returns:
            Background noise description
        """
        response_lower = response_text.lower()
        
        # Check for JSON-formatted noise info first
        try:
            json_match = re.search(r'"background_noise"\s*:\s*"([^"]+)"', response_lower)
            if json_match:
                return json_match.group(1)
        except:
            pass
        
        noise_indicators = ['none', 'static', 'hiss', 'hum', 'echo', 'reverb', 'music', 'background noise']
        for indicator in noise_indicators:
            if indicator in response_lower:
                return indicator
        return 'none detected'
    
    def _extract_language(self, response_text: str) -> str:
        """
        Extract language detection from Gemini response.
        
        Args:
            response_text: Raw response text from Gemini
            
        Returns:
            Detected language
        """
        response_lower = response_text.lower()
        
        # Check for JSON-formatted language info first
        try:
            json_match = re.search(r'"language_detected"\s*:\s*"([^"]+)"', response_lower)
            if json_match:
                return json_match.group(1)
        except:
            pass
        
        languages = ['english', 'spanish', 'french', 'german', 'chinese', 'japanese', 'portuguese', 'korean', 'arabic', 'hindi']
        for lang in languages:
            if lang in response_lower:
                return lang
        return 'unknown'
    
    def _calculate_speech_rate(self, segments: List[Dict]) -> float:
        """
        Calculate words per minute across all segments.
        
        Args:
            segments: List of transcript segments
            
        Returns:
            Words per minute (WPM)
        """
        if not segments:
            return 0.0
        
        total_words = sum(seg.get('word_count', 0) for seg in segments)
        if total_words == 0:
            return 0.0
        
        # Get time span
        timestamps = [seg.get('timestamp_seconds', 0) for seg in segments 
                     if seg.get('timestamp_seconds') is not None]
        if not timestamps or max(timestamps) == min(timestamps):
            return 0.0
        
        duration_minutes = (max(timestamps) - min(timestamps)) / 60.0
        if duration_minutes == 0:
            return 0.0
        
        return total_words / duration_minutes


# ============================================================================
# UNIFIED DOCUMENT PROCESSOR
# ============================================================================

class UnifiedDocumentProcessor:
    """
    Processes PDF, DOCX, and audio files with Gemini-powered analysis.
    """
    
    def __init__(self, api_key: str = None):
        """
        Initialize processor.
        
        Args:
            api_key: Google API key
        """
        self.api_key = api_key
        self.gemini = GeminiClient(api_key)
        
        logger.info("UnifiedDocumentProcessor initialized")
    
    def process_file(self, file_path: str, extract_images: bool = True) -> ProcessedDocument:
        """
        Process any supported file type.
        
        Args:
            file_path: Path to the file
            extract_images: Whether to extract and analyze images
            
        Returns:
            ProcessedDocument with all chunks
        """
        # Determine file type
        ext = os.path.splitext(file_path)[1].lower()
        
        # Generate document ID
        doc_id = self._generate_document_id(file_path)
        filename = os.path.basename(file_path)
        
        start_time = datetime.now()
        
        if ext == '.pdf':
            return self.process_pdf(file_path, doc_id, filename, extract_images)
        elif ext == '.docx':
            return self.process_docx(file_path, doc_id, filename, extract_images)
        elif ext in ['.mp3', '.wav', '.ogg', '.m4a', '.flac', '.aac', '.wma', '.aiff', 
                     '.opus', '.webm', '.amr', '.3gp', '.midi', '.mid', '.ra', '.ram', 
                     '.mp2', '.ac3', '.mp4', '.m4b', '.m4p', '.oga']:
            return self.process_audio(file_path, doc_id, filename)
        else:
            raise ValueError(f"Unsupported file type: {ext}")
    
    def process_pdf(
        self,
        file_path: str,
        doc_id: str,
        filename: str,
        extract_images: bool = True
    ) -> ProcessedDocument:
        """
        Process PDF with Gemini-powered analysis.
        
        Args:
            file_path: Path to PDF
            doc_id: Document ID
            filename: Original filename
            extract_images: Whether to extract images
            
        Returns:
            ProcessedDocument
        """
        if not PYMUPDF_AVAILABLE:
            raise ImportError("PyMuPDF not installed")
        
        start_time = datetime.now()
        chunks = []
        images = []
        tables = []
        
        try:
            doc = PyMuPDF.open(file_path)
            
            logger.info(f"Processing PDF: {filename} ({len(doc)} pages)")
            
            for page_num, page in enumerate(doc, 1):
                # 1. Extract text
                text = page.get_text()
                if text.strip():
                    text_chunks = self._chunk_text(text, page_num)
                    for i, chunk_text in enumerate(text_chunks):
                        chunk = DocumentChunk(
                            content=chunk_text,
                            chunk_type='text',
                            source_file=filename,
                            document_id=doc_id,
                            page_number=page_num,
                            metadata={'chunk_index': i, 'total_chunks': len(text_chunks)}
                        )
                        chunks.append(chunk)
                
                # 2. Extract tables
                page_tables = page.find_tables()
                if page_tables.tables:
                    for table_idx, table in enumerate(page_tables.tables):
                        # Convert to markdown
                        table_md = self._table_to_markdown(table)
                        
                        if table_md.strip():
                            # Analyze with Gemini
                            analysis = self.gemini.analyze_table(
                                table_md,
                                context=f"Page {page_num} of {filename}"
                            )
                            
                            table_chunk = DocumentChunk(
                                content=table_md,
                                chunk_type='table',
                                source_file=filename,
                                document_id=doc_id,
                                page_number=page_num,
                                table_data=table_md,
                                metadata={
                                    'table_index': table_idx,
                                    'analysis': analysis.get('summary', ''),
                                    'rows': len(table.rows),
                                    'columns': len(table.rows[0]) if table.rows else 0,
                                    # Enhanced metadata: structure
                                    'table_structure': analysis.get('structure', {}),
                                    'column_details': analysis.get('columns', []),
                                    'relationships': analysis.get('relationships', []),
                                    'patterns_insights': analysis.get('patterns_insights', []),
                                    'data_quality': analysis.get('data_quality', {}),
                                    'suggested_visualizations': analysis.get('suggested_visualizations', []),
                                    'potential_use_cases': analysis.get('potential_use_cases', []),
                                    'confidence_scores': analysis.get('confidence_scores', {}),
                                    # Convenience flags
                                    'has_numeric_data': any(col.get('data_type') in ['numeric', 'integer', 'float', 'currency', 'percentage']
                                                           for col in analysis.get('columns', [])),
                                    'has_date_data': any(col.get('data_type') in ['date', 'time']
                                                        for col in analysis.get('columns', [])),
                                    'key_columns': [col.get('column_index') for col in analysis.get('columns', [])
                                                   if col.get('is_key_column', False)],
                                    'column_count': len(analysis.get('columns', [])),
                                    'header_row_present': analysis.get('structure', {}).get('has_header_row', False)
                                }
                            )
                            chunks.append(table_chunk)
                            
                            tables.append({
                                'page': page_num,
                                'markdown': table_md,
                                'analysis': analysis
                            })
                
                # 3. Extract and analyze images
                if extract_images:
                    page_images = page.get_images()
                    for img_idx, img in enumerate(page_images):
                        try:
                            # Extract image bytes
                            xref = img[0]
                            pix = page.get_pixmap(matrix=PyMuPDF.Matrix(2, 2), clip=page.rect)
                            img_data = pix.tobytes('png')
                            
                            if len(img_data) > 500:  # Skip tiny images
                                # Analyze with Gemini Vision
                                analysis = self.gemini.analyze_image_with_vision(img_data)
                                
                                # Create image chunk with enhanced description
                                description = analysis.get('description', '[Image]')
                                keywords = analysis.get('search_keywords', analysis.get('keywords', []))
                                
                                # Build comprehensive content string including chart data
                                content_parts = [f"Image: {description}"]
                                if keywords:
                                    content_parts.append(f"Keywords: {', '.join(keywords)}")
                                
                                # Include chart/graph data in searchable content
                                data_viz = analysis.get('data_visualization', {})
                                if data_viz.get('chart_type', 'unknown') != 'unknown':
                                    content_parts.append(f"Chart type: {data_viz.get('chart_type')}")
                                    if data_viz.get('title'):
                                        content_parts.append(f"Chart title: {data_viz.get('title')}")
                                    if data_viz.get('trends'):
                                        content_parts.append(f"Trends: {', '.join(data_viz.get('trends', []))}")
                                
                                # Include OCR text in searchable content
                                text_content = analysis.get('text_content', {})
                                if text_content.get('raw_text'):
                                    content_parts.append(f"Text in image: {text_content.get('raw_text')}")
                                
                                img_content = '. '.join(content_parts)
                                
                                img_chunk = DocumentChunk(
                                    content=img_content,
                                    chunk_type='image',
                                    source_file=filename,
                                    document_id=doc_id,
                                    page_number=page_num,
                                    image_data=base64.b64encode(img_data).decode('utf-8'),
                                    metadata={
                                        'image_index': img_idx,
                                        'description': description,
                                        'keywords': keywords,
                                        'analysis': analysis.get('raw_response', ''),
                                        # Enhanced metadata: chart/graph interpretation
                                        'chart_type': data_viz.get('chart_type', 'unknown'),
                                        'text_blocks': text_content.get('text_blocks', []),
                                        'data_points': data_viz.get('data_points', []),
                                        'axes_info': data_viz.get('axes', {}),
                                        'legends_info': data_viz.get('legends', []),
                                        'chart_title': data_viz.get('title', ''),
                                        'chart_caption': data_viz.get('caption', ''),
                                        'trends': data_viz.get('trends', []),
                                        'statistics_visible': data_viz.get('statistics_visible', []),
                                        # Enhanced metadata: context
                                        'document_type': analysis.get('context', {}).get('document_type', 'unknown'),
                                        'subject_area': analysis.get('context', {}).get('subject_area', 'unknown'),
                                        'likely_purpose': analysis.get('context', {}).get('likely_purpose', 'unknown'),
                                        'page_layout': analysis.get('context', {}).get('page_layout', 'unknown'),
                                        # Enhanced metadata: object detection
                                        'objects_detected': analysis.get('objects_detected', []),
                                        # Enhanced metadata: confidence scores
                                        'confidence_scores': analysis.get('confidence_scores', {}),
                                        # Convenience flags
                                        'has_chart_data': len(data_viz.get('data_points', [])) > 0,
                                        'has_text_content': len(text_content.get('text_blocks', [])) > 0,
                                        'has_objects': len(analysis.get('objects_detected', [])) > 0
                                    }
                                )
                                chunks.append(img_chunk)
                                
                                images.append({
                                    'page': page_num,
                                    'index': img_idx,
                                    'data': img_data,
                                    'analysis': analysis
                                })
                                
                        except Exception as e:
                            logger.warning(f"Failed to extract image {img_idx}: {e}")
            
            doc.close()
            
            # Generate embeddings for all chunks
            logger.info(f"Generating embeddings for {len(chunks)} chunks...")
            self._generate_embeddings(chunks)
            
            processing_time = (datetime.now() - start_time).total_seconds()
            
            return ProcessedDocument(
                document_id=doc_id,
                filename=filename,
                file_type='pdf',
                chunks=chunks,
                images=images,
                tables=tables,
                processing_time=processing_time
            )
            
        except Exception as e:
            logger.error(f"PDF processing error: {e}")
            raise
    
    def process_docx(
        self,
        file_path: str,
        doc_id: str,
        filename: str,
        extract_images: bool = True
    ) -> ProcessedDocument:
        """
        Process Word document with Gemini-powered analysis.
        
        Args:
            file_path: Path to DOCX
            doc_id: Document ID
            filename: Original filename
            extract_images: Whether to extract images
            
        Returns:
            ProcessedDocument
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx not installed")
        
        start_time = datetime.now()
        chunks = []
        images = []
        tables = []
        
        try:
            doc = Document(file_path)
            
            logger.info(f"Processing DOCX: {filename}")
            
            # Process paragraphs (text content)
            full_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)
            
            # Chunk the text
            combined_text = '\n'.join(full_text)
            text_chunks = self._chunk_text(combined_text, page_number=1)
            
            for i, chunk_text in enumerate(text_chunks):
                chunk = DocumentChunk(
                    content=chunk_text,
                    chunk_type='text',
                    source_file=filename,
                    document_id=doc_id,
                    page_number=1,
                    metadata={'chunk_index': i}
                )
                chunks.append(chunk)
            
            # Process tables
            for table_idx, table in enumerate(doc.tables):
                table_md = self._docx_table_to_markdown(table)
                
                if table_md.strip():
                    analysis = self.gemini.analyze_table(
                        table_md,
                        context=f"Table {table_idx + 1} in {filename}"
                    )
                    
                    table_chunk = DocumentChunk(
                        content=table_md,
                        chunk_type='table',
                        source_file=filename,
                        document_id=doc_id,
                        page_number=1,
                        table_data=table_md,
                        metadata={
                            'table_index': table_idx,
                            'analysis': analysis.get('summary', ''),
                            'rows': len(table.rows),
                            'columns': len(table.columns),
                            # Enhanced metadata
                            'table_structure': analysis.get('structure', {}),
                            'column_details': analysis.get('columns', []),
                            'relationships': analysis.get('relationships', []),
                            'patterns_insights': analysis.get('patterns_insights', []),
                            'data_quality': analysis.get('data_quality', {}),
                            'suggested_visualizations': analysis.get('suggested_visualizations', []),
                            'potential_use_cases': analysis.get('potential_use_cases', []),
                            'confidence_scores': analysis.get('confidence_scores', {}),
                            'has_numeric_data': any(col.get('data_type') in ['numeric', 'integer', 'float', 'currency', 'percentage']
                                                   for col in analysis.get('columns', [])),
                            'has_date_data': any(col.get('data_type') in ['date', 'time']
                                                for col in analysis.get('columns', [])),
                            'key_columns': [col.get('column_index') for col in analysis.get('columns', [])
                                           if col.get('is_key_column', False)],
                            'column_count': len(analysis.get('columns', [])),
                            'header_row_present': analysis.get('structure', {}).get('has_header_row', False)
                        }
                    )
                    chunks.append(table_chunk)
                    
                    tables.append({
                        'index': table_idx,
                        'markdown': table_md,
                        'analysis': analysis
                    })
            
            # Process images (inline)
            if extract_images:
                # Get images from document
                for rel in doc.part.rels.values():
                    if "image" in rel.target_ref:
                        try:
                            img_bytes = rel.target_part.blob
                            
                            if len(img_bytes) > 500:
                                analysis = self.gemini.analyze_image_with_vision(img_bytes)
                                
                                description = analysis.get('description', '[Image]')
                                keywords = analysis.get('search_keywords', analysis.get('keywords', []))
                                
                                # Build comprehensive content string
                                content_parts = [f"Image: {description}"]
                                if keywords:
                                    content_parts.append(f"Keywords: {', '.join(keywords)}")
                                
                                data_viz = analysis.get('data_visualization', {})
                                if data_viz.get('chart_type', 'unknown') != 'unknown':
                                    content_parts.append(f"Chart type: {data_viz.get('chart_type')}")
                                    if data_viz.get('title'):
                                        content_parts.append(f"Chart title: {data_viz.get('title')}")
                                    if data_viz.get('trends'):
                                        content_parts.append(f"Trends: {', '.join(data_viz.get('trends', []))}")
                                
                                text_content = analysis.get('text_content', {})
                                if text_content.get('raw_text'):
                                    content_parts.append(f"Text in image: {text_content.get('raw_text')}")
                                
                                img_content = '. '.join(content_parts)
                                
                                img_chunk = DocumentChunk(
                                    content=img_content,
                                    chunk_type='image',
                                    source_file=filename,
                                    document_id=doc_id,
                                    page_number=1,
                                    image_data=base64.b64encode(img_bytes).decode('utf-8'),
                                    metadata={
                                        'image_index': len(images),
                                        'description': description,
                                        'keywords': keywords,
                                        'analysis': analysis.get('raw_response', ''),
                                        # Enhanced metadata
                                        'chart_type': data_viz.get('chart_type', 'unknown'),
                                        'text_blocks': text_content.get('text_blocks', []),
                                        'data_points': data_viz.get('data_points', []),
                                        'axes_info': data_viz.get('axes', {}),
                                        'legends_info': data_viz.get('legends', []),
                                        'chart_title': data_viz.get('title', ''),
                                        'chart_caption': data_viz.get('caption', ''),
                                        'trends': data_viz.get('trends', []),
                                        'statistics_visible': data_viz.get('statistics_visible', []),
                                        'document_type': analysis.get('context', {}).get('document_type', 'unknown'),
                                        'subject_area': analysis.get('context', {}).get('subject_area', 'unknown'),
                                        'likely_purpose': analysis.get('context', {}).get('likely_purpose', 'unknown'),
                                        'page_layout': analysis.get('context', {}).get('page_layout', 'unknown'),
                                        'objects_detected': analysis.get('objects_detected', []),
                                        'confidence_scores': analysis.get('confidence_scores', {}),
                                        'has_chart_data': len(data_viz.get('data_points', [])) > 0,
                                        'has_text_content': len(text_content.get('text_blocks', [])) > 0,
                                        'has_objects': len(analysis.get('objects_detected', [])) > 0
                                    }
                                )
                                chunks.append(img_chunk)
                                
                                images.append({
                                    'index': len(images),
                                    'data': img_bytes,
                                    'analysis': analysis
                                })
                        except Exception as e:
                            logger.warning(f"Failed to process image: {e}")
            
            # Generate embeddings
            logger.info(f"Generating embeddings for {len(chunks)} chunks...")
            self._generate_embeddings(chunks)
            
            processing_time = (datetime.now() - start_time).total_seconds()
            
            return ProcessedDocument(
                document_id=doc_id,
                filename=filename,
                file_type='docx',
                chunks=chunks,
                images=images,
                tables=tables,
                processing_time=processing_time
            )
            
        except Exception as e:
            logger.error(f"DOCX processing error: {e}")
            raise
    
    def process_audio(
        self,
        file_path: str,
        doc_id: str,
        filename: str
    ) -> ProcessedDocument:
        """
        Process audio file with transcription.
        
        Args:
            file_path: Path to audio file
            doc_id: Document ID
            filename: Original filename
            
        Returns:
            ProcessedDocument with transcript chunks
        """
        start_time = datetime.now()
        chunks = []
        
        try:
            # Read audio file
            with open(file_path, 'rb') as f:
                audio_data = f.read()
            
            # Determine MIME type
            ext = os.path.splitext(file_path)[1].lower()
            mime_type = {
                '.mp3': 'audio/mpeg',
                '.wav': 'audio/wav',
                '.ogg': 'audio/ogg',
                '.m4a': 'audio/mp4',
                '.flac': 'audio/flac',
                '.aac': 'audio/aac',
                '.wma': 'audio/x-ms-wma',
                '.aiff': 'audio/aiff',
                '.opus': 'audio/opus',
                '.webm': 'audio/webm',
                '.amr': 'audio/amr',
                '.3gp': 'audio/3gpp',
                '.midi': 'audio/midi',
                '.mid': 'audio/midi',
                '.ra': 'audio/x-pn-realaudio',
                '.ram': 'audio/x-pn-realaudio',
                '.mp2': 'audio/mpeg',
                '.ac3': 'audio/ac3',
                '.mp4': 'audio/mp4',
                '.m4b': 'audio/mp4',
                '.m4p': 'audio/mp4',
                '.oga': 'audio/ogg'
            }.get(ext, 'audio/mpeg')
            
            logger.info(f"Transcribing audio: {filename}")
            
            # Transcribe with Gemini
            result = self.gemini.transcribe_audio(audio_data, mime_type)
            
            if result['success']:
                # Create chunks for each segment with enhanced speaker metadata
                for segment in result['segments']:
                    timestamp = segment.get('timestamp_seconds', 0.0)
                    if timestamp == 0.0:
                        timestamp = self._parse_timestamp(segment['timestamp'])
                    
                    # Determine if this is a speaker change
                    is_speaker_change = (len(chunks) > 0 and 
                                        chunks[-1].metadata.get('speaker') != segment.get('speaker', 'Unknown'))
                    
                    chunk = DocumentChunk(
                        content=segment['text'],
                        chunk_type='audio',
                        source_file=filename,
                        document_id=doc_id,
                        timestamp=timestamp,
                        metadata={
                            'timestamp_str': segment['timestamp'],
                            'segment_index': len(chunks),
                            # Enhanced speaker diarization metadata
                            'speaker': segment.get('speaker', 'Unknown'),
                            'tone': segment.get('tone', 'neutral'),
                            'confidence': segment.get('confidence', 0.8),
                            'word_count': segment.get('word_count', 0),
                            'is_speaker_change': is_speaker_change
                        }
                    )
                    chunks.append(chunk)
                
                # Also create a full transcript chunk with speaker information
                unique_speakers = list(set(seg.get('speaker', 'Unknown') for seg in result.get('segments', [])))
                full_transcript = DocumentChunk(
                    content=result['full_text'],
                    chunk_type='audio',
                    source_file=filename,
                    document_id=doc_id,
                    metadata={
                        'is_full_transcript': True,
                        'speaker_segments': result.get('segments', []),
                        'unique_speakers': unique_speakers,
                        'diarization_success': len(result.get('segments', [])) > 1 and len(unique_speakers) > 1,
                        # Enhanced audio analysis metadata
                        'audio_quality': result.get('audio_quality', 'unknown'),
                        'background_noise': result.get('background_noise', 'none detected'),
                        'language_detected': result.get('language_detected', 'unknown'),
                        'estimated_speakers': result.get('estimated_speakers', 1),
                        'speech_rate_wpm': result.get('speech_rate_wpm', 0.0)
                    }
                )
                chunks.append(full_transcript)
                
            else:
                # Gemini failed - try Whisper fallback
                logger.warning(f"=== GEMINI FAILED, TRYING WHISPER FALLBACK for {filename} ===")
                logger.warning(f"Gemini error was: {result.get('error', 'Unknown')[:100]}")
                whisper_result = None
                try:
                    # Ensure ffmpeg is available in PATH for whisper
                    try:
                        import shutil as shutil_mod
                        from imageio_ffmpeg import get_ffmpeg_exe
                        ffmpeg_exe = get_ffmpeg_exe()
                        ffmpeg_dir = os.path.dirname(ffmpeg_exe)
                        
                        # Create a copy named ffmpeg.exe so whisper can find it
                        ffmpeg_standard = os.path.join(ffmpeg_dir, 'ffmpeg.exe')
                        if not os.path.exists(ffmpeg_standard):
                            shutil_mod.copy2(ffmpeg_exe, ffmpeg_standard)
                        
                        if ffmpeg_dir not in os.environ.get('PATH', ''):
                            os.environ['PATH'] = ffmpeg_dir + os.pathsep + os.environ.get('PATH', '')
                    except Exception as ffmpeg_err:
                        logger.warning(f"ffmpeg setup issue: {ffmpeg_err}")
                    
                    import whisper
                    logger.info(f"Gemini failed, falling back to Whisper: {filename}")
                    model = whisper.load_model("base")
                    whisper_transcription = model.transcribe(file_path)
                    whisper_text = whisper_transcription.get('text', '')
                    whisper_segments = whisper_transcription.get('segments', [])
                    
                    if whisper_text:
                        # Create chunks from Whisper segments
                        for seg in whisper_segments:
                            seg_start = seg.get('start', 0)
                            seg_text = seg.get('text', '').strip()
                            if seg_text:
                                chunk = DocumentChunk(
                                    content=seg_text,
                                    chunk_type='audio',
                                    source_file=filename,
                                    document_id=doc_id,
                                    timestamp=float(seg_start),
                                    metadata={
                                        'timestamp_str': f"[{int(seg_start // 60):02d}:{int(seg_start % 60):02d}]",
                                        'segment_index': len(chunks),
                                        'speaker': 'Speaker 1',
                                        'tone': 'neutral',
                                        'confidence': float(seg.get('confidence', 0.8)),
                                        'word_count': len(seg_text.split()),
                                        'is_speaker_change': False
                                    }
                                )
                                chunks.append(chunk)
                        
                        # Create full transcript chunk
                        full_transcript = DocumentChunk(
                            content=whisper_text,
                            chunk_type='audio',
                            source_file=filename,
                            document_id=doc_id,
                            metadata={
                                'is_full_transcript': True,
                                'unique_speakers': ['Speaker 1'],
                                'diarization_success': False,
                                'audio_quality': 'unknown',
                                'background_noise': 'unknown',
                                'language_detected': 'unknown',
                                'estimated_speakers': 1,
                                'speech_rate_wpm': 0.0,
                                'transcription_method': 'whisper'
                            }
                        )
                        chunks.append(full_transcript)
                        
                        # Update result for the return statement
                        result = {
                            'success': True,
                            'full_text': whisper_text,
                            'segments': []
                        }
                        logger.info(f"Whisper transcription succeeded: {len(chunks)} chunks")
                    else:
                        raise Exception("Whisper returned empty transcription")
                        
                except ImportError:
                    logger.warning("Whisper not installed")
                    chunk = DocumentChunk(
                        content=f"Transcription failed: Gemini quota exceeded. Install openai-whisper for offline transcription.",
                        chunk_type='audio',
                        source_file=filename,
                        document_id=doc_id,
                        metadata={'error': True}
                    )
                    chunks.append(chunk)
                    result = {'success': False, 'full_text': '', 'segments': []}
                    
                except Exception as whisper_error:
                    logger.error(f"Whisper fallback also failed: {whisper_error}")
                    chunk = DocumentChunk(
                        content=f"Transcription failed. Gemini: {result.get('error', 'Unknown error')[:100]}. Whisper error: {str(whisper_error)[:200]}",
                        chunk_type='audio',
                        source_file=filename,
                        document_id=doc_id,
                        metadata={'error': True}
                    )
                    chunks.append(chunk)
                    result = {'success': False, 'full_text': '', 'segments': []}
            
            # Generate embeddings
            logger.info(f"Generating embeddings for {len(chunks)} chunks...")
            self._generate_embeddings(chunks)
            
            processing_time = (datetime.now() - start_time).total_seconds()
            
            return ProcessedDocument(
                document_id=doc_id,
                filename=filename,
                file_type='audio',
                chunks=chunks,
                transcript=result.get('full_text', ''),
                processing_time=processing_time
            )
            
        except Exception as e:
            logger.error(f"Audio processing error: {e}")
            raise
    
    # ==================== HELPER METHODS ====================
    
    def _chunk_text(self, text: str, page_number: int = None, chunk_size: int = 1000) -> List[str]:
        """
        Split text into chunks.
        
        Args:
            text: Input text
            page_number: Page number for metadata
            chunk_size: Maximum chunk size
            
        Returns:
            List of text chunks
        """
        if len(text) <= chunk_size:
            return [text]
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + chunk_size
            
            # Try to break at sentence boundary
            if end < len(text):
                for punct in ['. ', '! ', '? ', '\n']:
                    last_punct = text[start:end].rfind(punct)
                    if last_punct != -1:
                        end = start + last_punct + 1
                        break
            
            chunks.append(text[start:end].strip())
            start = end - 200  # Overlap
        
        return [c for c in chunks if c.strip()]
    
    def _generate_embeddings(self, chunks: List[DocumentChunk]):
        """
        Generate Gemini Embedding 2.0 for all chunks.
        
        Args:
            chunks: List of DocumentChunk
        """
        if not chunks:
            return
        
        # Get texts to embed
        texts = [chunk.content for chunk in chunks]
        
        try:
            embeddings = self.gemini.embed_texts(texts)
            
            for chunk, embedding in zip(chunks, embeddings):
                chunk.embedding = embedding
                
        except Exception as e:
            logger.error(f"Embedding generation failed: {e}")
    
    def _table_to_markdown(self, table) -> str:
        """Convert PyMuPDF table to markdown."""
        if not table.rows:
            return ""
        
        lines = []
        
        # Header row
        header = [str(cell) if cell else '' for cell in table.rows[0]]
        lines.append('| ' + ' | '.join(header) + ' |')
        lines.append('| ' + ' | '.join(['---'] * len(header)) + ' |')
        
        # Data rows
        for row in table.rows[1:]:
            row_data = [str(cell) if cell else '' for cell in row]
            lines.append('| ' + ' | '.join(row_data) + ' |')
        
        return '\n'.join(lines)
    
    def _docx_table_to_markdown(self, table: Table) -> str:
        """Convert DOCX table to markdown."""
        if not table.rows:
            return ""
        
        lines = []
        
        # Header row
        header = [cell.text.strip() for cell in table.rows[0].cells]
        lines.append('| ' + ' | '.join(header) + ' |')
        lines.append('| ' + ' | '.join(['---'] * len(header)) + ' |')
        
        # Data rows
        for row in table.rows[1:]:
            row_data = [cell.text.strip() for cell in row.cells]
            lines.append('| ' + ' | '.join(row_data) + ' |')
        
        return '\n'.join(lines)
    
    def _parse_timestamp(self, timestamp_str: str) -> float:
        """Parse timestamp string to seconds."""
        try:
            parts = timestamp_str.split(':')
            if len(parts) == 2:
                return int(parts[0]) * 60 + int(parts[1])
            elif len(parts) == 3:
                return int(parts[0]) * 3600 + int(parts[1]) * 60 + int(parts[2])
        except:
            pass
        return 0.0
    
    def _generate_document_id(self, file_path: str) -> str:
        """Generate unique document ID."""
        timestamp = datetime.now().isoformat()
        hash_input = f"{file_path}_{timestamp}"
        return hashlib.sha256(hash_input.encode()).hexdigest()[:16]


# ============================================================================
# FACTORY FUNCTION
# ============================================================================

def create_processor(api_key: str = None) -> UnifiedDocumentProcessor:
    """
    Create unified document processor.
    
    Args:
        api_key: Google API key
        
    Returns:
        UnifiedDocumentProcessor instance
    """
    return UnifiedDocumentProcessor(api_key)


# ============================================================================
# MAIN TEST
# ============================================================================

if __name__ == "__main__":
    print("=" * 60)
    print("Unified Document Processor - Test")
    print("=" * 60)
    
    api_key = os.environ.get('GOOGLE_API_KEY')
    
    if not api_key:
        print("\nNo GOOGLE_API_KEY found!")
        print("Set GOOGLE_API_KEY to test")
    else:
        print("\nAPI key found, testing...")
        
        try:
            processor = UnifiedDocumentProcessor(api_key)
            print("Processor created successfully!")
            
            # Test embedding
            test_texts = ["Hello world", "Machine learning"]
            embeddings = processor.gemini.embed_texts(test_texts)
            print(f"Embedding test: {len(embeddings)} embeddings generated")
            print(f"Dimension: {len(embeddings[0])}")
            
        except Exception as e:
            print(f"Error: {e}")
    
    print("\n" + "=" * 60)
